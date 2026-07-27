"""Text processor for MD, DOCX, HTML, and plain text files."""

from __future__ import annotations

import re
from pathlib import Path
from typing import TYPE_CHECKING

import structlog
from pydantic import ValidationError

from course_supporter.ingestion.base import (
    CategorisedProcessingError,
    MaterialProcessor,
    UnsupportedFormatError,
)
from course_supporter.ingestion.paragraph_anchors import compute_paragraph_anchors
from course_supporter.ingestion.schemas import (
    DocumentSegmentDraft,
    DocumentSummaryDraft,
)
from course_supporter.language import display_name
from course_supporter.llm.error_categories import StructuralRetryError
from course_supporter.models.source import (
    ChunkType,
    ContentChunk,
    SourceDocument,
    SourceType,
)
from course_supporter.security.exceptions import ErrorCategory

if TYPE_CHECKING:
    from course_supporter.llm.stage_router import StageRouter
    from course_supporter.storage.orm import AuthoredDocument

logger = structlog.get_logger()

SUPPORTED_EXTENSIONS = {".md", ".markdown", ".docx", ".html", ".htm", ".txt"}

# Chunk types that count as paragraphs for the Phase 3.3a anchor (HEADING
# chunks are skipped — they do not advance the paragraph ordinal).
_TEXT_PARAGRAPH_TYPES = frozenset({ChunkType.PARAGRAPH})


class TextProcessor(MaterialProcessor):
    """Process text documents (MD, DOCX, HTML, TXT).

    Extracts headings and paragraphs without LLM.
    """

    async def process_raw(
        self,
        source: AuthoredDocument,
    ) -> SourceDocument:
        if source.source_type != SourceType.TEXT:
            raise UnsupportedFormatError(
                f"TextProcessor expects 'text', got '{source.source_type}'"
            )

        path = Path(source.source_url)
        ext = path.suffix.lower()

        if ext not in SUPPORTED_EXTENSIONS:
            raise UnsupportedFormatError(
                f"Unsupported text format: {ext}. "
                f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )

        logger.info(
            "text_processing_start",
            source_url=source.source_url,
            format=ext,
        )

        if ext in {".md", ".markdown"}:
            chunks = self._process_markdown(path)
        elif ext == ".docx":
            chunks = self._process_docx(path)
        elif ext in {".html", ".htm"}:
            chunks = self._process_html(path)
        else:  # .txt
            chunks = self._process_plain_text(path)

        logger.info(
            "text_processing_done",
            source_url=source.source_url,
            chunk_count=len(chunks),
        )

        return SourceDocument(
            source_type=SourceType.TEXT,
            source_url=source.source_url,
            title=source.filename or path.stem,
            chunks=chunks,
        )

    def _process_markdown(self, path: Path) -> list[ContentChunk]:
        """Parse Markdown file into heading/paragraph chunks.

        Headings: lines starting with # (1-6 levels).
        Paragraphs: non-empty text between headings.
        """
        content = path.read_text(encoding="utf-8")
        return self._parse_markdown_text(content)

    @staticmethod
    def _parse_markdown_text(content: str) -> list[ContentChunk]:
        """Parse markdown text content into chunks."""
        chunks: list[ContentChunk] = []
        heading_pattern = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)

        idx = 0
        last_end = 0

        for match in heading_pattern.finditer(content):
            # Text before this heading → paragraph
            before = content[last_end : match.start()].strip()
            if before:
                chunks.append(
                    ContentChunk(
                        chunk_type=ChunkType.PARAGRAPH,
                        text=before,
                        index=idx,
                    )
                )
                idx += 1

            # Heading itself
            level = len(match.group(1))
            heading_text = match.group(2).strip()
            chunks.append(
                ContentChunk(
                    chunk_type=ChunkType.HEADING,
                    text=heading_text,
                    index=idx,
                    metadata={"level": level},
                )
            )
            idx += 1
            last_end = match.end()

        # Remaining text after last heading
        remaining = content[last_end:].strip()
        if remaining:
            chunks.append(
                ContentChunk(
                    chunk_type=ChunkType.PARAGRAPH,
                    text=remaining,
                    index=idx,
                )
            )

        return chunks

    @staticmethod
    def _process_docx(path: Path) -> list[ContentChunk]:
        """Parse DOCX file into heading/paragraph chunks.

        Detects heading styles (Heading 1..6) and maps to levels.
        """
        from docx import Document

        doc = Document(str(path))
        chunks: list[ContentChunk] = []
        idx = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""

            if style_name.startswith("Heading"):
                # Extract level from "Heading 1", "Heading 2", etc.
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1

                chunks.append(
                    ContentChunk(
                        chunk_type=ChunkType.HEADING,
                        text=text,
                        index=idx,
                        metadata={"level": level},
                    )
                )
            else:
                chunks.append(
                    ContentChunk(
                        chunk_type=ChunkType.PARAGRAPH,
                        text=text,
                        index=idx,
                    )
                )
            idx += 1

        return chunks

    @staticmethod
    def _process_html(path: Path) -> list[ContentChunk]:
        """Parse HTML file into heading/paragraph chunks.

        Uses BeautifulSoup to extract <h1>..<h6> and <p> elements.
        """
        from bs4 import BeautifulSoup

        html = path.read_text(encoding="utf-8")
        soup = BeautifulSoup(html, "html.parser")
        chunks: list[ContentChunk] = []
        idx = 0

        for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p"]):
            text = element.get_text(strip=True)
            if not text:
                continue

            if element.name.startswith("h"):
                level = int(element.name[1])
                chunks.append(
                    ContentChunk(
                        chunk_type=ChunkType.HEADING,
                        text=text,
                        index=idx,
                        metadata={"level": level},
                    )
                )
            else:
                chunks.append(
                    ContentChunk(
                        chunk_type=ChunkType.PARAGRAPH,
                        text=text,
                        index=idx,
                    )
                )
            idx += 1

        return chunks

    @staticmethod
    def _process_plain_text(path: Path) -> list[ContentChunk]:
        """Read plain text file as a single paragraph chunk."""
        content = path.read_text(encoding="utf-8").strip()
        if not content:
            return []
        return [
            ContentChunk(
                chunk_type=ChunkType.PARAGRAPH,
                text=content,
                index=0,
            )
        ]

    async def process_macro(
        self,
        doc: SourceDocument,
        router: StageRouter,
    ) -> DocumentSummaryDraft:
        """Pass 2a -- premium LLM extracts concept structure (KD-2.1-A).

        Concatenates chunk text into a single body and routes the
        resulting document through the ``pass_2a_mapping`` stage.
        Returns a :class:`DocumentSummaryDraft`; segment drafts are
        retained on the result for Phase 2.1 C7 (Pass 2b) consumption
        and are not materialised here.

        Document-level ``main_concepts`` / ``secondary_concepts`` are
        derived algorithmically as ``union + dedup`` over the per-segment
        concepts (vision.md §2.2, KD-2.1-O). The LLM is asked to emit
        concepts only at segment level — this method assembles the
        document-level view by sorted set-union and applies the
        conflict rule: any concept that appears as ``main`` in at
        least one segment stays in ``main_concepts`` and is removed
        from ``secondary_concepts``.
        """
        text = doc.assemble_text()
        if not text.strip():
            msg = "Cannot run Pass 2a on empty document (no content chunks)"
            raise CategorisedProcessingError(ErrorCategory.EMPTY_DOCUMENT, msg)
        reference_text_length = len(text)
        parsed: dict[str, DocumentSummaryDraft] = {}

        def _coverage_validator(content: str) -> None:
            """StageRouter response_validator hook (fixup 2.1.7.2).

            Translates a Pydantic ``ValidationError`` into
            :class:`StructuralRetryError` so the router's existing
            instructor-style retry path fires (one retry on the
            same model with feedback appended; ladder fallback on
            second failure).
            """
            try:
                draft_local = DocumentSummaryDraft.model_validate_json(
                    content,
                    context={"reference_text_length": reference_text_length},
                )
            except ValidationError as exc:
                error_types = sorted({e.get("type", "unknown") for e in exc.errors()})
                first = exc.errors()[0]
                first_loc = ".".join(str(x) for x in first.get("loc", []))
                logger.warning(
                    "pass2a.validation.failed",
                    source_type=doc.source_type.value,
                    validation_error_types=error_types,
                    first_error_msg=first.get("msg", ""),
                    first_error_loc=first_loc,
                    reference_text_length=reference_text_length,
                )
                feedback = (
                    f"{first.get('msg', 'validation error')} "
                    f"(field: {first_loc or '<root>'}). "
                    "Regenerate the response with valid output."
                )
                raise StructuralRetryError(feedback) from exc
            parsed["draft"] = draft_local

        result = await router.execute_for_stage(
            "pass_2a_mapping",
            response_validator=_coverage_validator,
            expects_json=True,
            text=text,
            language=display_name(doc.language) if doc.language else None,
        )
        # Validator stored the parsed draft on the winning attempt;
        # if execute_for_stage returned, validator succeeded at least
        # once (otherwise it would have raised LadderExhaustedError).
        draft = parsed["draft"]
        logger.debug(
            "pass2a.validation.ok",
            source_type=doc.source_type.value,
            provider=result.provider_used,
            model=result.model_used,
            attempt_count=result.attempt_count,
            reference_text_length=reference_text_length,
        )

        # Algorithmic aggregation of document-level concepts from
        # per-segment concepts (vision.md §2.2, KD-2.1-O). LLM emits
        # concepts only at segment level; ``DocumentSummary.main_concepts``
        # is the sorted set-union over all segments. Conflict rule:
        # any concept that appears as ``main`` in at least one segment
        # stays in ``main_concepts`` and is removed from
        # ``secondary_concepts``.
        all_main: set[str] = set()
        all_secondary: set[str] = set()
        for seg in draft.segments:
            all_main.update(seg.main_concepts)
            all_secondary.update(seg.secondary_concepts)
        all_secondary -= all_main
        draft.main_concepts = sorted(all_main)
        draft.secondary_concepts = sorted(all_secondary)

        return draft

    async def process_detail(
        self,
        doc: SourceDocument,
        summary_draft: DocumentSummaryDraft,
    ) -> list[DocumentSegmentDraft]:
        """Pass 2b -- algorithmic slice over Pass 2a offsets (KD-2.1-O).

        Zero LLM calls. For each ``DocumentSegmentDraft`` in
        ``summary_draft.segments`` (Pass 2a output), materialises
        ``content`` by slicing ``doc.assemble_text()`` with the
        ``start_pos`` / ``end_pos`` offsets the mapping LLM emitted.

        Reference text is assembled via the canonical
        :meth:`SourceDocument.assemble_text` helper -- identical string
        seen by Pass 2a's mapping prompt and by Stage 2's safety check.

        Drafts that already carry a non-None ``content`` are passed through
        verbatim (defensive: not expected in text/web canonical path).
        Offset bounds are validated by ``DocumentSegmentRepository.create_batch``
        downstream so a single source of truth handles boundary errors.
        """
        reference_text = doc.assemble_text()
        # Paragraph anchors (Phase 3.3a): the emitted chunks line up with
        # assemble_text; PARAGRAPH chunks carry the ordinal, HEADING chunks
        # are skipped.
        emitted = [c for c in doc.chunks if c.text]
        filled: list[DocumentSegmentDraft] = []
        for draft in summary_draft.segments:
            start_para, end_para = compute_paragraph_anchors(
                emitted, _TEXT_PARAGRAPH_TYPES, draft.start_pos, draft.end_pos
            )
            update: dict[str, object] = {
                "start_paragraph": start_para,
                "end_paragraph": end_para,
            }
            if draft.content is None:
                update["content"] = reference_text[draft.start_pos : draft.end_pos]
            filled.append(draft.model_copy(update=update))
        return filled
